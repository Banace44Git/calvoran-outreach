#!/usr/bin/env python3
"""c5 — Brief-Ansprache: personalisierte Anschreiben je Lead als .docx.

Liest eine Auswahl-CSV (Spalte `name`, optional `website`), joint gegen
`calvoran.companies` (Adresse + GF) und `calvoran.dossiers` (geschaeftsmodell),
erzeugt je Lead mit Sonnet (`ansprache_saetze`, models.yaml) drei variable Bausteine
und merged sie in die Word-Vorlage (JTILS-v3.docx). Layout/QR der Vorlage bleiben 1:1.

Variable Slots (Anker = JTILS-Beispieltext in der Vorlage):
  - Adressblock: Herrn/Frau, Vorname Nachname, Firma, Straße Nr., PLZ Ort
  - Anrede:      "Sehr geehrter Herr Jacob,"
  - Hook:        erster Satz des Absatzes "Sie führen mit JTILS …"; Rest wird aus
                 der Vorlage abgeleitet (folgt v3-Edits).
  - Bullet 1:    "der Optimierung des Warenlagers"
  - Bullet 2:    "Vertriebscontrolling sowie Auftragscontrolling und -kalkulation"

Aufruf:
  .venv/bin/python pipeline/c5_brief_merge.py \
      --selection <auswahl.csv> --template <JTILS-v3.docx> --outdir <briefe-dir>

Idempotent re Inhalte: bereits abgenommene Sätze aus <outdir>/_merge-data.json
(oder --reuse <json>) werden je Lead wiederverwendet, sonst frisch generiert.
"""
from __future__ import annotations

import argparse
import csv
import json
import os
import re
import sys
import unicodedata

import anthropic
import docx
import gender_guesser.detector as gg

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from calvoran.db import get_client  # noqa: E402

MODEL = "claude-sonnet-4-6"  # == config/models.yaml tasks.ansprache_saetze.primary
HR_GF_DEFAULT = ("/Users/johannesbreuers/projects/os/01-projects/fractional-cfo/"
                 "hr-abruf/gf-geburtsdaten.csv")
ALT_FLAG = 75   # GF ab diesem Alter zur bewussten Bestätigung markieren

# Anker in der Vorlage (Platzhalter bzw. JTILS-Beispieltext)
A_ANREDE_BLOCK = "Herrn/Frau"
A_NAME = "Vorname Nachname"
A_FIRMA = "Firma"
A_STRASSE = "Straße Nr."
A_PLZ_ORT = "PLZ Ort"
A_SALUT = "Sehr geehrter Herr Jacob,"
A_HOOK_PREFIX = "Sie führen mit JTILS"
A_HOOK_SPLIT = "Eine Frage stellt sich"            # ab hier ist der Absatz fix
A_BULLET1 = "der Optimierung des Warenlagers"
A_BULLET2 = "Vertriebscontrolling sowie Auftragscontrolling und -kalkulation"

# GF-Sonderfälle, die aus dem Rohnamen nicht eindeutig parsebar sind
GF_OVERRIDE = {  # name-substring -> (vorname, nachname, gender) — übersteuert HR-Daten
    "JTILS": ("Mathew", "Jacob", "m"),       # Register: Mathew; Jo adressiert "Herr Jacob" (m.jacob@)
    "Günter Wendt": ("Frank", "Wendt", "m"),  # ältester (Therese 86) = Platzhalterdatum/Seniorin; Jo -> Frank (59)
    # GF ohne Geburtsdatum im AD (Alt-Eintrag): nicht in gf-geburtsdaten.csv erfasst,
    # einziger datierter Eintrag war ein Prokurist (jetzt korrekt ist_gf=0). GF-Name direkt
    # aus dem AD ("Geschäftsführer: <Nachname>, <Vorname>, ...").
    "Iser GmbH": ("Joachim", "Iser", "m"),
    "Schilles": ("Peter", "Schilles", "m"),       # H + P Schilles Tiefbau-GmbH
    "Heerdt": ("Hans-Peter", "Heerdt", "m"),       # Raumausstattung Heerdt GmbH
}
GENDER_OVERRIDE = {  # name-substring -> 'm'/'f', wenn gender-guesser unschlüssig (Jo bestätigt)
    "MFT Membran-Filtrations": "f",      # Hongmei Yan
}

SYS = ("Du textest Brief-Bausteine für Johannes Breuers, externer CFO (Verkaufsvorbereitung im Mittelstand). "
       "Senior-Praktiker-Deutsch, knapp, ganze Sätze, generisches Maskulinum. Keine Marketing-Floskeln, "
       "keine Gedankenstriche, keine LLM-Floskeln. Erde dich AUSSCHLIESSLICH am gelieferten Geschäftsmodell; "
       "erfinde keine Fakten. Antworte NUR mit JSON.")


def gen_prompt(name: str, wz: str, geschaeftsmodell: str) -> str:
    return (f"Firma: {name}\nBranche (WZ): {wz}\nGeschäftsmodell: {geschaeftsmodell}\n\n"
            "Erzeuge JSON mit drei Feldern:\n"
            '1. "hook": EXAKT Muster: "Sie führen mit <Kurzname> ein gut aufgestelltes Unternehmen und kennen '
            'den Wert von <X>." <Kurzname>=natürliche Kurzform/Markenname (nicht volle Rechtsform); '
            '<X>=Kernkompetenz aus dem Geschäftsmodell, grammatisch passend zu "den Wert von <X>".\n'
            '2. "beispiel1": Dativ-/Nominalphrase als Ersatz für Listenpunkt "der Optimierung des Warenlagers". '
            'Betriebsnah, firmenspezifisch, passend nach "Konkret helfe ich zum Beispiel bei: ". '
            'Kurz, kein Schlusspunkt, kein "Ich helfe bei".\n'
            '3. "beispiel2": ZWEITER, thematisch ANDERER firmenspezifischer Listenpunkt als Ersatz für '
            '"Vertriebscontrolling sowie Auftragscontrolling und -kalkulation". Gleiche Formregeln; nicht mit '
            'beispiel1 überschneiden; eher Controlling-/Kalkulations-/Steuerungs-nah.')


_GENDER = gg.Detector(case_sensitive=False)
_TITLES = ["Prof.", "Dr.", "Dipl.-Ing.", "Dipl.", "Ing."]


def _norm(s):
    s = unicodedata.normalize("NFKD", (s or "").lower()).encode("ascii", "ignore").decode()
    for t in [" gmbh", " mbh", " kg", " co", " ohg", " ag", "."]:
        s = s.replace(t, " ")
    return " ".join(s.split())


def _split_title(vor, nach):
    title = " ".join(t for t in _TITLES if t in f"{vor} {nach}")
    for t in _TITLES:
        vor, nach = vor.replace(t, ""), nach.replace(t, "")
    return " ".join(vor.split()), " ".join(nach.split()), title.strip()


def load_oldest_gf(hr_path):
    """norm(firma) -> (alter, vorname, nachname, title) des ältesten aktiven GF."""
    groups = {}
    with open(hr_path, newline="", encoding="utf-8") as f:
        for r in csv.DictReader(f):
            if r.get("ist_gf") != "1":
                continue
            try:
                alter = int(r["gf_alter"])
            except (ValueError, TypeError):
                continue
            groups.setdefault(_norm(r["firma"]), []).append(
                (alter, r["gf_vorname"].strip(), r["gf_nachname"].strip()))
    out = {}
    for k, lst in groups.items():
        alter, vor, nach = max(lst, key=lambda x: x[0])
        vor, nach, title = _split_title(vor, nach)
        out[k] = (alter, vor, nach, title)
    return out


def _gender(name, vorname):
    for key, gd in GENDER_OVERRIDE.items():
        if key in name:
            return gd
    g = _GENDER.get_gender(vorname.split()[0].split("-")[0]) if vorname else "unknown"
    return "m" if g in ("male", "mostly_male") else ("f" if g in ("female", "mostly_female") else "?")


def parse_gf(name, oldest_map):
    """Anrede für genau einen GF — den ältesten. -> (adr1, name_zeile, salut, flags)."""
    flags = []
    for key, (vor, nach, gd) in GF_OVERRIDE.items():
        if key in name:
            adr = "Herrn" if gd == "m" else "Frau"
            return adr, f"{vor} {nach}", f"Sehr geehrte{'r' if gd=='m' else ''} {'Herr' if gd=='m' else 'Frau'} {nach},", flags
    rec = oldest_map.get(_norm(name))
    if not rec:
        return None, "", "Sehr geehrte Damen und Herren,", ["KEINE GF-Altersdaten — generische Anrede"]
    alter, vor, nach, title = rec
    gd = _gender(name, vor)
    if gd == "?":
        gd = "m"
        flags.append(f"Geschlecht unklar (Vorname '{vor}')")
    if alter >= ALT_FLAG:
        flags.append(f"ältester GF {alter} J. — bewusst ansprechen?")
    adr = "Herrn" if gd == "m" else "Frau"
    tpref = f"{title} " if title else ""
    salut = f"Sehr geehrte{'r' if gd=='m' else ''} {'Herr' if gd=='m' else 'Frau'} {tpref}{nach},"
    return adr, f"{tpref}{vor} {nach}".strip(), salut, flags


def generate(client, name, wz, gm):
    msg = client.messages.create(model=MODEL, max_tokens=500, temperature=0, system=SYS,
                                 messages=[{"role": "user", "content": gen_prompt(name, wz, gm)}])
    txt = re.sub(r"^```(json)?|```$", "", msg.content[0].text.strip(), flags=re.M).strip()
    o = json.loads(txt)
    return o["hook"].strip(), o["beispiel1"].strip().rstrip("."), o["beispiel2"].strip().rstrip(".")


def hook_rest(template_path):
    """Fixer Rest des Hook-Absatzes aus der Vorlage (ab 'Eine Frage stellt sich')."""
    for p in docx.Document(template_path).paragraphs:
        if p.text.strip().startswith(A_HOOK_PREFIX):
            idx = p.text.find(A_HOOK_SPLIT)
            if idx >= 0:
                return p.text[idx:].strip()
    raise SystemExit("Hook-Anker nicht in Vorlage gefunden — Template-Drift?")


def merge(template, outpath, *, adr1, name_zeile, firma, strasse, plz_ort, salut, hook, rest, b1, b2):
    d = docx.Document(template)
    repl = {A_ANREDE_BLOCK: adr1, A_NAME: name_zeile, A_FIRMA: firma, A_STRASSE: strasse,
            A_PLZ_ORT: plz_ort, A_SALUT: salut, A_BULLET1: b1, A_BULLET2: b2}
    seen = set()
    for p in d.paragraphs:
        t = p.text.strip()
        if not p.runs:
            continue
        if t in repl:
            new = repl[t]; seen.add(t)
        elif t.startswith(A_HOOK_PREFIX):
            new = f"{hook} {rest}"; seen.add("HOOK")
        else:
            continue
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    d.save(outpath)
    return seen


def slug(s):
    return re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")[:40]


def record_outreach(cl, company_ids, wave):
    """Idempotent je Brief eine outreach-Zeile (channel='letter', status='queued') anlegen.
    Fetch-existing-then-insert (keine Dubletten, läuft auch ohne Unique-Index aus 0006).
    Rückgabe: Anzahl neu angelegter Zeilen."""
    seen = set()
    for i in range(0, len(company_ids), 50):
        rows = (cl.table("outreach").select("company_id")
                .eq("channel", "letter").eq("wave", wave)
                .in_("company_id", company_ids[i:i + 50]).execute().data)
        seen.update(r["company_id"] for r in rows)
    new = [{"company_id": cid, "channel": "letter", "status": "queued", "wave": wave}
           for cid in company_ids if cid not in seen]
    for i in range(0, len(new), 50):
        cl.table("outreach").insert(new[i:i + 50]).execute()
    return len(new)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--selection", required=True, help="CSV mit Spalte 'name'")
    ap.add_argument("--template", required=True, help="Word-Vorlage (JTILS-v3.docx)")
    ap.add_argument("--outdir", required=True)
    ap.add_argument("--reuse", help="merge-data.json mit abgenommenen Sätzen (key=name-substring)")
    ap.add_argument("--gf-data", default=HR_GF_DEFAULT, help="hr-abruf GF-Geburtsdaten CSV (ältester GF)")
    ap.add_argument("--wave", type=int, help="Versandwelle: legt je Brief eine outreach-Zeile "
                                             "(status='queued') für die CRM-Nachverfolgung an")
    args = ap.parse_args()
    os.makedirs(args.outdir, exist_ok=True)
    rest = hook_rest(args.template)
    oldest_gf = load_oldest_gf(args.gf_data)

    # Auswahl laden
    names, wz_csv = [], {}
    with open(args.selection, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            if str(row.get("Lead", "TRUE")).upper() in ("TRUE", "1", "JA"):
                names.append(row["name"]); wz_csv[row["name"]] = row.get("branche_wz", "")

    # Abgenommene Sätze (Wiederverwendung)
    reuse = {}
    reuse_path = args.reuse or os.path.join(os.path.dirname(args.outdir), "briefe-test-2026-06-25", "_merge-data.json")
    if os.path.exists(reuse_path):
        reuse = json.load(open(reuse_path))

    cl = get_client()
    comp = {}
    for i in range(0, len(names), 50):
        for c in cl.table("companies").select(
                "id,name,strasse,plz,ort,ges_vertreter,anzahl_gf,branche_wz").in_("name", names[i:i + 50]).execute().data:
            comp[c["name"]] = c
    ids = [c["id"] for c in comp.values()]
    doss = {}
    for i in range(0, len(ids), 50):
        for d in cl.table("dossiers").select("company_id,dossier").in_("company_id", ids[i:i + 50]).execute().data:
            doss[d["company_id"]] = d["dossier"] or {}

    client = anthropic.Anthropic()
    merge_data, review, merged_ids = {}, [], []
    for name in names:
        c = comp.get(name)
        if not c:
            review.append({"name": name, "flags": ["KEIN companies-Match — übersprungen"]}); continue
        dd = doss.get(c["id"], {})
        # Sätze: abgenommen wiederverwenden, sonst generieren
        appr = next((v for k, v in reuse.items() if k in name), None)
        if appr:
            hook, b1, b2, src = appr["hook"], appr["b1"], appr["b2"], "abgenommen"
        else:
            hook, b1, b2 = generate(client, name, c.get("branche_wz") or wz_csv.get(name, ""),
                                    dd.get("geschaeftsmodell") or ""); src = "sonnet"
        adr1, name_zeile, salut, flags = parse_gf(name, oldest_gf)
        merge(args.template, os.path.join(args.outdir, f"{slug(name)}.docx"),
              adr1=adr1, name_zeile=name_zeile, firma=name, strasse=c.get("strasse") or "",
              plz_ort=f"{c.get('plz') or ''} {c.get('ort') or ''}".strip(), salut=salut,
              hook=hook, rest=rest, b1=b1, b2=b2)
        merge_data[name] = {"hook": hook, "b1": b1, "b2": b2}
        merged_ids.append(c["id"])
        review.append({"name": name, "anrede": salut, "src": src, "hook": hook, "b1": b1, "b2": b2, "flags": flags})

    json.dump(merge_data, open(os.path.join(args.outdir, "_merge-data.json"), "w"), ensure_ascii=False, indent=2)
    _write_review(os.path.join(args.outdir, "_review.md"), review)
    n_ok = sum(1 for r in review if "hook" in r)
    n_flag = sum(1 for r in review if r.get("flags"))
    print(f"{n_ok}/{len(names)} Briefe erzeugt -> {args.outdir} | {n_flag} mit Flags (siehe _review.md)")

    if args.wave is not None and merged_ids:
        n_new = record_outreach(cl, merged_ids, args.wave)
        print(f"outreach (Welle {args.wave}): {n_new} neue 'queued'-Zeilen, "
              f"{len(merged_ids) - n_new} bereits vorhanden.")


def _write_review(path, review):
    flagged = [r for r in review if r.get("flags")]
    ok = [r for r in review if not r.get("flags")]
    L = ["# Brief-Review — c5 Ansprache", "",
         f"{len(review)} Leads · {len(flagged)} zur Kontrolle markiert.", ""]
    if flagged:
        L += ["## Zur Kontrolle (Anrede/GF prüfen)", ""]
        for r in flagged:
            L.append(f"### {r['name']}")
            L.append(f"- Flags: {'; '.join(r['flags'])}")
            if "anrede" in r:
                L += [f"- Anrede: {r['anrede']}", f"- Hook: {r['hook']}",
                      f"- B1: {r['b1']}", f"- B2: {r['b2']}"]
            L.append("")
    L += ["## OK", "", "| Firma | Anrede | Hook (…Wert von) | B1 | B2 |", "|---|---|---|---|---|"]
    for r in ok:
        wert = r["hook"].split("kennen den Wert von", 1)[-1].strip().rstrip(".")
        L.append(f"| {r['name']} | {r['anrede']} | {wert} | {r['b1']} | {r['b2']} |")
    open(path, "w", encoding="utf-8").write("\n".join(L) + "\n")


if __name__ == "__main__":
    main()
